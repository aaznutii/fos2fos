import os

import docx2txt as docx2txt

from full_copy import deepcopy

from docx import Document as Doc
from docx.parts import image as parts_img
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.document import Document
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.shared import Pt, Cm

PATH_fos = 'Docs/fos/'
PATH_fos_result = 'Docs/fos_result/'

doc = Doc(f'{PATH_fos}ФОС_Феномен факта в современных масс-медиа_ДО.docx')
output_doc = Doc()


# Функция итерации по всем блокам документа для получения таблиц и параграфов
def iter_block_items(parent):
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def create_docx_file():
    for block in iter_block_items(doc):
        # Работа с параграфом
        try:
            par = block.text
            par = output_doc.add_paragraph(par)
            # установить стили для объектов параграфа
            for run in par.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

        # Работа с таблицей
        except AttributeError:
            new_tbl = deepcopy(block._tbl)
            paragraph = output_doc.add_paragraph()
            # After that, we add the previously copied table
            paragraph._p.addnext(new_tbl)

    output_doc.save('Docs/fos_result/output3.docx')
    # Пока не удается только забрать картинку


def get_img(doc_path):
    # Extract the images to img_folder/
    docx2txt.process(doc_path, 'Docs/img_folder/')
    # Save all 'rId:filenames' relationships in an dictionary named rels
    docs = Doc(doc_path)
    # blip = docs.
    rels = {}
    for r in docs.part.rels.values():
        if isinstance(r._target, parts_img.ImagePart):
            rels[r.rId] = os.path.basename(r._target.partname)
    print(rels)

    # Then process your text
    for paragraph in doc.paragraphs:
        # If you find an image
        if 'Graphic' in paragraph._p.xml:
            # Get the rId of the image
            for rId in rels:
                if rId in paragraph._p.xml:
                    path = os.path.join('Docs/img_folder/', rels[rId])
                    p = output_doc.add_paragraph()
                    p = p.add_run()
                    p.add_picture(path)
                    output_doc.save('Docs/fos_result/img.docx')
        else:
            # It's not an image
            continue


# get_img(f'{PATH_fos}ФОС_Феномен факта в современных масс-медиа_ДО.docx')
create_docx_file()

